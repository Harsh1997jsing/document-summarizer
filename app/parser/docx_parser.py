from docx import Document
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file using python-docx.

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        str: Extracted text from paragraphs and tables.
    """
    text_parts = []

    try:
        doc = Document(file_path)
        logger.info(f"Opened DOCX: {file_path}")


        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        if not text_parts:
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return ""

        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return "\n".join(text_parts)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX '{file_path}': {e}")
        raise RuntimeError(f"Failed to parse DOCX: {file_path}") from e